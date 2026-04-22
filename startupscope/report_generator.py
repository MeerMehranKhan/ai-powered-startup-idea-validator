import os
import json
import csv
from datetime import datetime
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
from startupscope.models import ValidationReport, ExportMetadata
from startupscope.logger import logger
from startupscope.exceptions import ReportGenerationError, InvalidExportFormatError

def generate_export_filename(report: ValidationReport, ext: str) -> str:
    safe_name = "".join([c if c.isalnum() else "_" for c in report.input_data.name])
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M")
    return f"StartupScope_Report_{safe_name}_{timestamp}.{ext}"

def export_to_json(report: ValidationReport, filepath: str):
    with open(filepath, 'w', encoding='utf-8') as f:
        f.write(report.model_dump_json(indent=4))

def export_to_csv(report: ValidationReport, filepath: str):
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(["Startup Name", "Score", "Risk Level", "Recommendation"])
        writer.writerow([
            report.input_data.name, 
            report.scores.overall_score, 
            report.scores.risk_level, 
            report.executive_summary.recommendation
        ])

def export_to_pdf(report: ValidationReport, filepath: str):
    try:
        c = canvas.Canvas(filepath, pagesize=letter)
        c.drawString(100, 750, f"StartupScope AI Report: {report.input_data.name}")
        c.drawString(100, 730, f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        c.drawString(100, 710, f"Overall Score: {report.scores.overall_score}/100")
        c.drawString(100, 690, f"Risk Level: {report.scores.risk_level}")
        
        c.drawString(100, 650, "Executive Summary:")
        c.drawString(100, 630, report.executive_summary.short_summary[:100] + "...")
        c.drawString(100, 610, f"Recommendation: {report.executive_summary.recommendation}")
        c.save()
    except Exception as e:
        logger.error(f"PDF generation failed: {e}")
        raise ReportGenerationError("Failed to generate PDF document.")

def export_to_docx(report: ValidationReport, filepath: str):
    try:
        doc = Document()
        doc.add_heading(f"StartupScope AI Report: {report.input_data.name}", 0)
        doc.add_paragraph(f"Overall Score: {report.scores.overall_score}/100")
        doc.add_paragraph(f"Risk Level: {report.scores.risk_level}")
        
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(report.executive_summary.short_summary)
        doc.add_paragraph(f"Recommendation: {report.executive_summary.recommendation}")
        
        doc.save(filepath)
    except Exception as e:
        logger.error(f"DOCX generation failed: {e}")
        raise ReportGenerationError("Failed to generate DOCX document.")

def generate_report_file(report: ValidationReport, format_type: str = "pdf") -> ExportMetadata:
    format_type = format_type.lower()
    if format_type not in ["pdf", "docx", "json", "csv"]:
        raise InvalidExportFormatError(f"Unsupported export format: {format_type}")
        
    filename = generate_export_filename(report, format_type)
    folder_path = os.path.join("exports", format_type)
    os.makedirs(folder_path, exist_ok=True)
    
    filepath = os.path.join(folder_path, filename)
    
    try:
        if format_type == "pdf":
            export_to_pdf(report, filepath)
        elif format_type == "docx":
            export_to_docx(report, filepath)
        elif format_type == "json":
            export_to_json(report, filepath)
        elif format_type == "csv":
            export_to_csv(report, filepath)
            
        logger.info(f"Successfully generated {format_type.upper()} report: {filepath}")
        return ExportMetadata(
            filename=filepath,
            format=format_type.upper(),
            timestamp=datetime.now().isoformat(),
            status="Success"
        )
    except ReportGenerationError:
        logger.warning(f"Failed to generate {format_type.upper()}. Falling back to JSON.")
        fallback_filepath = os.path.join("exports", "json", generate_export_filename(report, "json"))
        export_to_json(report, fallback_filepath)
        return ExportMetadata(
            filename=fallback_filepath,
            format="JSON (Fallback)",
            timestamp=datetime.now().isoformat(),
            status="Fallback Success"
        )
